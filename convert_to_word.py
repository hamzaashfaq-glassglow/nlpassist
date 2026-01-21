"""
Convert technical report markdown to Word document
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Create document
doc = Document()

# Add title
title = doc.add_heading('NLPAssist+ Technical Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('A Retrieval-Augmented Generation System for NLP Question Answering')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.italic = True
subtitle_format.font.size = Pt(12)

doc.add_page_break()

# Read markdown file
with open(r'C:\Users\Supreme_Traders\.gemini\antigravity\brain\7dcc9e36-a145-4773-9cf7-87d2c121deb1\technical_report.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Remove title and subtitle (already added)
content = re.sub(r'^# NLPAssist\+ Technical Report.*?\n---\n', '', content, flags=re.DOTALL)

# Process content line by line
lines = content.split('\n')
in_code_block = False
code_lang = None
in_table = False
table_headers = []

i = 0
while i < len(lines):
    line = lines[i]
    
    # Skip mermaid diagrams (can't render in Word easily)
    if line.strip().startswith('```mermaid'):
        # Skip until end of code block
        i += 1
        while i < len(lines) and not lines[i].strip().startswith('```'):
            i += 1
        i += 1
        doc.add_paragraph('[Architecture Diagram - See Markdown File]', style='Intense Quote')
        continue
    
    # Code blocks
    if line.strip().startswith('```'):
        if not in_code_block:
            in_code_block = True
            code_lang = line.strip()[3:]
            i += 1
            continue
        else:
            in_code_block = False
            i += 1
            continue
    
    if in_code_block:
        p = doc.add_paragraph(line, style='No Spacing')
        p.runs[0].font.name = 'Courier New'
        p.runs[0].font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.5)
        i += 1
        continue
    
    # Headings
    if line.startswith('# '):
        doc.add_heading(line[2:], 1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], 2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], 3)
    elif line.startswith('#### '):
        doc.add_heading(line[5:], 4)
    
    # Tables
    elif '|' in line and line.strip().startswith('|'):
        if not in_table:
            # Start of table
            in_table = True
            table_headers = [cell.strip() for cell in line.split('|')[1:-1]]
            # Skip separator line
            i += 2
            
            # Count rows
            row_count = 0
            j = i
            while j < len(lines) and lines[j].strip().startswith('|'):
                row_count += 1
                j += 1
            
            # Create table
            table = doc.add_table(rows=row_count + 1, cols=len(table_headers))
            table.style = 'Light Grid Accent 1'
            
            # Add headers
            for col_idx, header in enumerate(table_headers):
                cell = table.rows[0].cells[col_idx]
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
            
            # Add rows
            row_idx = 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                for col_idx, cell_text in enumerate(cells):
                    if col_idx < len(table_headers):
                        table.rows[row_idx].cells[col_idx].text = cell_text
                i += 1
                row_idx += 1
            
            in_table = False
            doc.add_paragraph()  # Add spacing after table
            continue
    
    # Bold and italic
    elif line.strip():
        p = doc.add_paragraph()
        
        # Process inline formatting
        text = line
        
        # Bold
        while '**' in text:
            before, rest = text.split('**', 1)
            if '**' in rest:
                bold_text, after = rest.split('**', 1)
                p.add_run(before)
                p.add_run(bold_text).bold = True
                text = after
            else:
                p.add_run(before)
                text = '**' + rest
                break
        
        p.add_run(text)
    
    # Empty line
    else:
        if i > 0 and lines[i-1].strip():  # Only add space if previous line had content
            doc.add_paragraph()
    
    i += 1

# Save document
output_path = r'C:\Users\Supreme_Traders\Desktop\nlpassist\NLPAssist_Technical_Report.docx'
doc.save(output_path)
print(f"✅ Document saved to: {output_path}")
